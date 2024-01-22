from docx.shared import Inches
import docx
from docx2pdf import convert
import pandas as pd
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication 
from email.mime.multipart import MIMEMultipart 
import smtplib 
import os 

# private variables the public does not need
from app import app_pass, filepath, from_email

def message(subject="2023 Tax donation receipt Dragonettes",  
            text="", img=None, 
            attachment=None): 
    
    # build message contents 
    msg = MIMEMultipart() 
      
    # Add Subject 
    msg['Subject'] = subject   
      
    # Add text contents 
    msg.attach(MIMEText(text)) 

    if attachment is not None:

        if type(attachment) is not list: 
            
                # if it isn't a list, make it one 
            attachment = [attachment]   

        for one_attachment in attachment: 

            with open(one_attachment, 'rb') as f: 
                
                # Read in the attachment 
                # using MIMEApplication 
                file = MIMEApplication( 
                    f.read(), 
                    name=os.path.basename(one_attachment) 
                ) 
            file['Content-Disposition'] = f'attachment; filename="{os.path.basename(one_attachment)}"'
            
            msg.attach(file)

    return msg

# cannot have these / - : in sponsor name

#### Variables to set ####
df = pd.read_csv("filename.csv")
convert_docx_to_pdf = 1
send_email = 1
#### ####


if __name__ == "__main__":
    # update tax year in email_body_base and p1

    if send_email:
        email_body_base = """,

    We gratefully appreciate your donation to the Round Rock High School Dragonette Booster Club during the 2023-2024 school year.  Your tax receipt is attached.

    Your generous donation will be used for competition fees, dance apparel, choreography, props, and other operating expenses not covered by the dance department’s district budget. In addition, donations also help the Booster Club provide scholarships to graduating seniors. Thank you for partnering with us to make this year so special for the dancers and their school community.

    Jason Schwarz
    Sponsorship Committee Co-chair
    RRHS Dragonette Booster Club"""

        smtp = smtplib.SMTP('smtp.gmail.com', 587) 
        smtp.ehlo() 
        smtp.starttls()
        smtp.login(from_email, app_pass)

    start = 0
    stop = len(df)

    for idx in range(start, stop):
        # pull info from dataframe
        sponsor = df['sponsor'][idx]
        contact = df['contact'][idx]
        date = df['date'][idx]
        donation = df['donation'][idx]
        donation_type = df['donation_type'][idx]
        email_address = df['email'][idx]

        # create document and add to it
        document = docx.Document()

        ##### document starts with logo #####
        document.add_picture('dragonette_logo.png')

        # title and heading
        document.add_heading("Dragonette Donation Document", level=0)

        document.add_paragraph("Round Rock High School Dragonette Booster Club").paragraph_format.space_after = Inches(0)
        document.add_paragraph("110 N IH 35, Suite 315-174").paragraph_format.space_after = Inches(0)
        document.add_paragraph("Round Rock, TX 78681")

        document.add_paragraph("Tax ID: 47-0873690").paragraph_format.space_after = Inches(0.5)


        # body 
        p1 = document.add_paragraph("The Round Rock High School Dragonette Booster Club, a ")
        p1.add_run("501(c)(3) organization").bold = True
        p1.add_run(", thanks you for your generous donation below for the 2023-2024 Dragonettes. Donations may be tax-deductible depending on your tax situation.")

        p2a = document.add_paragraph("Sponsor: ")
        p2a.add_run(sponsor).bold = True

        p2b = document.add_paragraph("Contact: ")
        p2b.add_run(contact).bold = True

        p3a = document.add_paragraph("Donation: ")
        p3a.add_run(donation).bold = True

        p3b = document.add_paragraph("Donation type: ")
        p3b.add_run(donation_type).bold = True

        p4 = document.add_paragraph("Donation date: ")
        p4.add_run(date).bold = True
        p4.paragraph_format.space_after = Inches(0.5)


        # signature
        document.add_paragraph("Appreciatively,").paragraph_format.space_after = Inches(0)
        document.add_picture("jason_schwarz_signature.png")  # insert your own signature file here
        document.add_paragraph("Jason Schwarz").paragraph_format.space_after = Inches(0)
        document.add_paragraph("Sponsorship Committee Co-chair").paragraph_format.space_after = Inches(0)
        document.add_paragraph("Round Rock High School Dragonette Booster Club")
        ##### document ends with signature #####

        # save file
        filename = "2023 Tax donation receipt " + sponsor + ".docx"
        document.save(filename)

        # convert to pdf
        if convert_docx_to_pdf:
            convert(filename)

        if send_email:
            # setup email specifics
            email_body = "Dear " + contact + email_body_base
            to_email = [email_address] 
            
            # generate pdf file path
            pdf_filename = filename[:-4] + "pdf"
            att_file_path = filepath + pdf_filename
            
            # create email message 
            msg = message(text=email_body, attachment=att_file_path) 

            # send the email
            smtp.sendmail(from_addr=from_email,
                        to_addrs=to_email, 
                        msg=msg.as_string()) 

    if send_email:
        smtp.quit()